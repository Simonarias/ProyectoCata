import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import zipfile
from datetime import datetime
import io
import tempfile

# Configuración de la página
st.set_page_config(
    page_title="Generador de Contratos Automático",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

def aplicar_formato_texto(paragraph, texto, es_variable=False):
    """
    Aplica formato Arial 12 y resalta variables en azul
    """
    paragraph.clear()
    run = paragraph.add_run(texto)
    
    # Aplicar formato Arial 12
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    
    # Si es una variable, aplicar color azul
    if es_variable:
        run.font.color.rgb = RGBColor(0, 112, 192)  # Azul
        run.bold = True

def reemplazar_en_paragraph(paragraph, variables):
    """
    Reemplaza variables en un párrafo manteniendo formato
    """
    texto_original = paragraph.text
    texto_nuevo = texto_original
    
    # Reemplazar variables
    for variable, valor in variables.items():
        if variable in texto_nuevo:
            texto_nuevo = texto_nuevo.replace(variable, valor)
    
    # Si hubo cambios, aplicar formato
    if texto_original != texto_nuevo:
        paragraph.clear()
        
        # Dividir texto para aplicar formato diferente a variables
        texto_partes = texto_nuevo
        for variable, valor in variables.items():
            if valor in texto_partes:
                partes = texto_partes.split(valor)
                paragraph.clear()
                
                for i, parte in enumerate(partes):
                    if i > 0:
                        # Agregar la variable con formato especial
                        run_var = paragraph.add_run(valor)
                        run_var.font.name = 'Arial'
                        run_var.font.size = Pt(12)
                        run_var.font.color.rgb = RGBColor(0, 112, 192)  # Azul
                        run_var.bold = True
                    
                    if parte:
                        # Agregar texto normal
                        run_normal = paragraph.add_run(parte)
                        run_normal.font.name = 'Arial'
                        run_normal.font.size = Pt(12)
                break
        else:
            # Si no hay variables, solo aplicar formato normal
            run = paragraph.add_run(texto_nuevo)
            run.font.name = 'Arial'
            run.font.size = Pt(12)

def generar_contratos(df, plantilla_bytes):
    """
    Genera contratos automáticamente desde una plantilla Word y datos de DataFrame
    """
    contratos_generados = []
    errores = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for index, row in df.iterrows():
        try:
            status_text.text(f"Procesando contrato {index + 1} de {len(df)}: {row['NOMBRE']}")
            
            # Crear documento desde la plantilla
            doc = Document(io.BytesIO(plantilla_bytes))
            
            # Diccionario con las variables y sus valores
            variables = {
                "[(CONTRATO NÚMERO)]": str(row["CONTRATO NÚMERO"]) if pd.notna(row["CONTRATO NÚMERO"]) else "",
                "[(CÉDULA)]": str(row["CÉDULA"]) if pd.notna(row["CÉDULA"]) else "",
                "[(CORREO ELECTRÓNICO)]": str(row["CORREO ELECTRÓNICO"]) if pd.notna(row["CORREO ELECTRÓNICO"]) else "",
                "[(NOMBRE)]": str(row["NOMBRE"]) if pd.notna(row["NOMBRE"]) else "",
                "[(FECHA DE INICIO)]": str(row["FECHA DE INICIO"]) if pd.notna(row["FECHA DE INICIO"]) else "",
                "[(FECHA FINALIZACIÓN)]": str(row["FECHA FINALIZACIÓN"]) if pd.notna(row["FECHA FINALIZACIÓN"]) else "",
                "[(PLAZO EN DÍAS)]": str(row["PLAZO EN DÍAS"]) if pd.notna(row["PLAZO EN DÍAS"]) else "",
                "[(VALOR TOTAL DEL CONTRATO SIN IVA)]": str(row["VALOR TOTAL DEL CONTRATO SIN IVA"]) if pd.notna(row["VALOR TOTAL DEL CONTRATO SIN IVA"]) else "",
            }
            
            # Reemplazar variables en párrafos
            for paragraph in doc.paragraphs:
                reemplazar_en_paragraph(paragraph, variables)
            
            # Reemplazar variables en tablas
            for table in doc.tables:
                for row_table in table.rows:
                    for cell in row_table.cells:
                        for paragraph in cell.paragraphs:
                            reemplazar_en_paragraph(paragraph, variables)
            
            # Reemplazar variables en headers y footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        reemplazar_en_paragraph(paragraph, variables)
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        reemplazar_en_paragraph(paragraph, variables)
            
            # Generar nombre del archivo
            cedula = str(row["CÉDULA"]).replace("/", "-").replace("\\", "-")
            nombre = str(row["NOMBRE"]).replace("/", "-").replace("\\", "-")
            contrato = str(row["CONTRATO NÚMERO"]).replace("/", "-").replace("\\", "-")
            nombre_archivo = f"Contrato_{nombre}_{cedula}_{contrato}.docx"
            
            # Guardar en memoria
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            contratos_generados.append({
                'nombre': nombre_archivo,
                'contenido': doc_bytes.getvalue(),
                'info': f"{row['NOMBRE']} - {row['CÉDULA']}"
            })
            
        except Exception as e:
            errores.append(f"Error en fila {index + 1} ({row.get('NOMBRE', 'N/A')}): {str(e)}")
        
        # Actualizar progreso
        progress_bar.progress((index + 1) / len(df))
    
    status_text.text("¡Proceso completado!")
    return contratos_generados, errores

def crear_zip_contratos(contratos):
    """
    Crea un archivo ZIP con todos los contratos
    """
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for contrato in contratos:
            zip_file.writestr(contrato['nombre'], contrato['contenido'])
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def main():
    # Título principal
    st.title("📄 Generador de Contratos Automático")
    st.markdown("---")
    
    # Sidebar con información
    with st.sidebar:
        st.header("ℹ️ Información")
        st.info("""
        **Instrucciones:**
        1. Sube tu plantilla Word (.docx)
        2. Sube tu archivo Excel (.xlsx)
        3. Verifica los datos en vista previa
        4. Genera los contratos
        5. Descarga el archivo ZIP
        """)
        
        st.header("📋 Variables requeridas")
        variables_info = [
            "[(CONTRATO NÚMERO)]",
            "[(CÉDULA)]", 
            "[(CORREO ELECTRÓNICO)]",
            "[(NOMBRE)]",
            "[(FECHA DE INICIO)]",
            "[(FECHA FINALIZACIÓN)]",
            "[(PLAZO EN DÍAS)]",
            "[(VALOR TOTAL DEL CONTRATO SIN IVA)]"
        ]
        
        for var in variables_info:
            st.code(var, language="text")
    
    # Columnas principales
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📁 Archivos de entrada")
        
        # Upload plantilla Word
        plantilla_file = st.file_uploader(
            "Sube tu plantilla Word (.docx)",
            type=['docx'],
            help="Archivo Word con las variables a reemplazar"
        )
        
        # Upload Excel
        excel_file = st.file_uploader(
            "Sube tu archivo Excel (.xlsx)",
            type=['xlsx'],
            help="Archivo Excel con los datos de los contratos"
        )
    
    with col2:
        st.header("👁️ Vista previa de datos")
        
        if excel_file is not None:
            try:
                df = pd.read_excel(excel_file)
                st.success(f"✅ Excel cargado: {len(df)} registros encontrados")
                
                # Mostrar columnas
                st.subheader("Columnas encontradas:")
                columnas_requeridas = [
                    "CONTRATO NÚMERO", "CÉDULA", "CORREO ELECTRÓNICO", "NOMBRE",
                    "FECHA DE INICIO", "FECHA FINALIZACIÓN", "PLAZO EN DÍAS", 
                    "VALOR TOTAL DEL CONTRATO SIN IVA"
                ]
                
                for col in columnas_requeridas:
                    if col in df.columns:
                        st.success(f"✅ {col}")
                    else:
                        st.error(f"❌ {col} - FALTANTE")
                
                # Vista previa de datos
                st.subheader("Vista previa:")
                st.dataframe(df.head(), use_container_width=True)
                
            except Exception as e:
                st.error(f"❌ Error al leer Excel: {str(e)}")
        else:
            st.info("📤 Sube un archivo Excel para ver la vista previa")
    
    # Sección de generación
    st.markdown("---")
    st.header("🚀 Generación de contratos")
    
    if plantilla_file is not None and excel_file is not None:
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col2:
            if st.button("🔄 Generar Contratos", type="primary", use_container_width=True):
                try:
                    # Leer archivos
                    df = pd.read_excel(excel_file)
                    plantilla_bytes = plantilla_file.read()
                    
                    # Validar columnas
                    columnas_requeridas = [
                        "CONTRATO NÚMERO", "CÉDULA", "CORREO ELECTRÓNICO", "NOMBRE",
                        "FECHA DE INICIO", "FECHA FINALIZACIÓN", "PLAZO EN DÍAS", 
                        "VALOR TOTAL DEL CONTRATO SIN IVA"
                    ]
                    
                    columnas_faltantes = [col for col in columnas_requeridas if col not in df.columns]
                    
                    if columnas_faltantes:
                        st.error(f"❌ Faltan columnas: {', '.join(columnas_faltantes)}")
                    else:
                        # Generar contratos
                        with st.spinner("Generando contratos..."):
                            contratos, errores = generar_contratos(df, plantilla_bytes)
                        
                        # Mostrar resultados
                        if contratos:
                            st.success(f"🎉 ¡{len(contratos)} contratos generados exitosamente!")
                            
                            # Crear ZIP
                            zip_data = crear_zip_contratos(contratos)
                            
                            # Botón de descarga
                            st.download_button(
                                label="📥 Descargar todos los contratos (ZIP)",
                                data=zip_data,
                                file_name=f"contratos_generados_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                                mime="application/zip",
                                type="primary",
                                use_container_width=True
                            )
                            
                            # Mostrar lista de contratos generados
                            st.subheader("📋 Contratos generados:")
                            for i, contrato in enumerate(contratos, 1):
                                st.write(f"{i}. {contrato['info']}")
                        
                        # Mostrar errores si los hay
                        if errores:
                            st.error("⚠️ Se encontraron algunos errores:")
                            for error in errores:
                                st.write(f"• {error}")
                
                except Exception as e:
                    st.error(f"❌ Error general: {str(e)}")
    else:
        st.info("📤 Sube ambos archivos (plantilla Word y Excel) para continuar")
    
    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: #666;'>
            💡 Las variables reemplazadas aparecerán resaltadas en <span style='color: #0070C0; font-weight: bold;'>azul</span> en los documentos generados
        </div>
        """, 
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()